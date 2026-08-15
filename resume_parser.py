"""
resume_parser.py
Extracts plain text and a lightly-structured view of an uploaded resume
(PDF, DOCX, or TXT) so the rest of the pipeline can work with it.
"""
import os
import re
import docx
import pdfplumber
from logger_config import setup_logger

# Set up logger for this module
logger = setup_logger(__name__)

SECTION_HEADERS = [
    "summary", "professional summary", "objective", "profile",
    "skills", "core competencies", "technical skills",
    "experience", "work experience", "professional experience", "employment history",
    "education", "certifications", "certification", "projects",
    "achievements", "awards", "publications", "languages", "interests",
]


def extract_text(file_path: str) -> str:
    logger.info(f"Starting text extraction from file: {file_path}")
    ext = os.path.splitext(file_path)[1].lower()
    logger.info(f"Detected file extension: {ext}")
    
    if ext == ".pdf":
        logger.info("Processing PDF file")
        return _extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        logger.info("Processing DOCX/DOC file")
        return _extract_docx(file_path)
    elif ext == ".txt":
        logger.info("Processing TXT file")
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
            logger.info(f"Successfully extracted {len(text)} characters from TXT file")
            return text
    else:
        logger.error(f"Unsupported file type: {ext}")
        raise ValueError(f"Unsupported resume file type: {ext}")


def _extract_pdf(file_path: str) -> str:
    logger.info(f"Starting PDF extraction from: {file_path}")
    text_chunks = []
    try:
        with pdfplumber.open(file_path) as pdf:
            logger.info(f"PDF has {len(pdf.pages)} pages")
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""
                if page_text:
                    text_chunks.append(page_text)
                    logger.debug(f"Extracted {len(page_text)} characters from page {i+1}")
                else:
                    logger.warning(f"No text extracted from page {i+1}")
        
        result = "\n".join(text_chunks)
        logger.info(f"PDF extraction completed. Total characters: {len(result)}")
        return result
    except Exception as e:
        logger.error(f"Error extracting PDF: {e}", exc_info=True)
        raise


def _extract_docx(file_path: str) -> str:
    logger.info(f"Starting DOCX extraction from: {file_path}")
    try:
        d = docx.Document(file_path)
        logger.info(f"Document has {len(d.paragraphs)} paragraphs and {len(d.tables)} tables")
        
        parts = []
        for para in d.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        
        logger.info(f"Extracted {len(parts)} text paragraphs")
        
        for table_idx, table in enumerate(d.tables):
            logger.debug(f"Processing table {table_idx+1}")
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        
        result = "\n".join(parts)
        logger.info(f"DOCX extraction completed. Total characters: {len(result)}")
        return result
    except Exception as e:
        logger.error(f"Error extracting DOCX: {e}", exc_info=True)
        raise


def split_sections(raw_text: str) -> dict:
    """
    Very lightweight section splitter: scans line by line, treats a line that
    matches (case-insensitively, allowing punctuation/whitespace) one of the
    known section headers as a new section boundary.
    """
    logger.info("Starting section splitting")
    logger.debug(f"Input text length: {len(raw_text)} characters")
    
    lines = [l.strip() for l in raw_text.splitlines()]
    logger.info(f"Processing {len(lines)} lines")
    
    sections = {"header": []}
    current = "header"
    header_matches = 0

    for line in lines:
        if not line:
            continue
        normalized = re.sub(r"[^a-z ]", "", line.lower()).strip()
        matched_header = None
        for header in SECTION_HEADERS:
            if normalized == header or (len(normalized) < 40 and normalized.startswith(header) and len(normalized.split()) <= 4):
                matched_header = header
                break
        if matched_header:
            current = matched_header
            sections.setdefault(current, [])
            header_matches += 1
            logger.debug(f"Found section header: {matched_header}")
            continue
        sections.setdefault(current, []).append(line)

    result = {k: "\n".join(v).strip() for k, v in sections.items() if "\n".join(v).strip()}
    logger.info(f"Section splitting completed. Found {len(result)} sections: {list(result.keys())}")
    logger.info(f"Total section headers matched: {header_matches}")
    return result


def extract_contact_info(raw_text: str) -> dict:
    logger.info("Starting contact information extraction")
    logger.debug(f"Input text length: {len(raw_text)} characters")
    
    email_match = re.search(r"[\w\.\-+]+@[\w\-]+\.[\w\.\-]+", raw_text)
    phone_match = re.search(r"(\+?\d{1,3}[\s\-]?)?(\(?\d{2,4}\)?[\s\-]?){2,4}\d{3,4}", raw_text)
    linkedin_match = re.search(r"(linkedin\.com/[^\s,]+)", raw_text, re.IGNORECASE)
    
    email = email_match.group(0) if email_match else ""
    phone = phone_match.group(0).strip() if phone_match else ""
    linkedin = linkedin_match.group(0) if linkedin_match else ""
    
    logger.info(f"Email found: {bool(email)}")
    logger.info(f"Phone found: {bool(phone)}")
    logger.info(f"LinkedIn found: {bool(linkedin)}")
    
    name = ""
    for line in raw_text.splitlines():
        stripped = line.strip()
        if stripped and len(stripped.split()) <= 5 and not re.search(r"[@\d]", stripped):
            name = stripped
            logger.info(f"Potential name found: {name}")
            break
    
    if not name:
        logger.warning("No name could be extracted from resume")
    
    contact_info = {
        "name": name,
        "email": email,
        "phone": phone,
        "linkedin": linkedin,
    }
    logger.info(f"Contact extraction completed: {contact_info}")
    return contact_info
