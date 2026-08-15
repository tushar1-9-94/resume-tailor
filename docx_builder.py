"""
docx_builder.py
Renders the tailored resume data into a clean, ATS-friendly .docx:
single column, no tables/text boxes/images, standard fonts, standard
section headers, reverse-chronological layout.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from logger_config import setup_logger

# Set up logger for this module
logger = setup_logger(__name__)


FONT_NAME = "Calibri"


def _set_font(run, size=11, bold=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    rFonts = run._element.rPr.rFonts
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if color:
        run.font.color.rgb = color
    logger.debug(f"Set font: {FONT_NAME}, size: {size}, bold: {bold}")


def _add_section_spacing(paragraph, before=12, after=6):
    """Add consistent spacing to paragraphs"""
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.15


def _add_heading(doc, text):
    logger.debug(f"Adding heading: {text}")
    p = doc.add_paragraph()
    _add_section_spacing(p, before=18, after=6)
    run = p.add_run(text.upper())
    _set_font(run, size=13, bold=True)
    # Add border for section divider
    pPr = p._p.get_or_add_pPr()
    from docx.oxml import OxmlElement
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "666666")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_contact_paragraph(doc, text, size=10, bold=False, center=False):
    """Add a contact information paragraph with proper formatting"""
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_section_spacing(p, before=2, after=2)
    run = p.add_run(text)
    _set_font(run, size=size, bold=bold)
    return p


def build_resume_docx(output_path: str, contact: dict, tailored: dict, job_title: str = ""):
    logger.info("=== Starting DOCX resume building ===")
    logger.info(f"Output path: {output_path}")
    logger.debug(f"Contact info: {contact}")
    logger.debug(f"Job title: {job_title}")
    logger.debug(f"Tailored data keys: {list(tailored.keys())}")
    
    doc = Document()
    logger.info("Created new Document instance")

    # Set up document margins (standard resume margins)
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    logger.debug("Set document margins to standard resume dimensions")

    # Set default document font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    logger.debug(f"Set default font to {FONT_NAME} 11pt")

    # --- Header Section ---
    logger.info("Building header section")
    
    # Name (larger, bold, centered)
    name = contact.get("name") or "Your Name"
    logger.info(f"Adding name to header: {name}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_section_spacing(p, before=0, after=4)
    run = p.add_run(name)
    _set_font(run, size=16, bold=True)

    # Contact information (centered, smaller)
    contact_bits = [b for b in [contact.get("email"), contact.get("phone"), contact.get("linkedin")] if b]
    logger.info(f"Adding {len(contact_bits)} contact information items")
    if contact_bits:
        _add_contact_paragraph(doc, " | ".join(contact_bits), size=10, center=True)

    # Target role (if available)
    if job_title:
        logger.info(f"Adding target role: {job_title}")
        _add_contact_paragraph(doc, f"Target Role: {job_title}", size=10, bold=True, center=True)

    # Add horizontal line after header
    p = doc.add_paragraph()
    _add_section_spacing(p, before=8, after=12)
    pPr = p._p.get_or_add_pPr()
    from docx.oxml import OxmlElement
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # --- Professional Summary ---
    if tailored.get("professional_summary"):
        logger.info("Adding Professional Summary section")
        _add_heading(doc, "Professional Summary")
        p = doc.add_paragraph()
        _add_section_spacing(p, before=6, after=6)
        run = p.add_run(tailored["professional_summary"])
        _set_font(run, size=11)

    # --- Core Skills ---
    skills = tailored.get("core_skills") or []
    if skills:
        logger.info(f"Adding Core Skills section with {len(skills)} skills")
        _add_heading(doc, "Core Skills")
        # Format skills as a clean, readable list
        p = doc.add_paragraph()
        _add_section_spacing(p, before=6, after=6)
        # Split skills into multiple lines if too many
        skills_text = " | ".join(skills)
        run = p.add_run(skills_text)
        _set_font(run, size=10)

    # --- Professional Experience ---
    experience = tailored.get("experience") or []
    if experience:
        logger.info(f"Adding Professional Experience section with {len(experience)} entries")
        _add_heading(doc, "Professional Experience")
        
        for idx, entry in enumerate(experience):
            logger.debug(f"Processing experience entry {idx+1}/{len(experience)}")
            
            # Add spacing between experience entries
            if idx > 0:
                p = doc.add_paragraph()
                _add_section_spacing(p, before=6, after=0)
            
            # Job title and company (bold, larger)
            title = entry.get("title", "")
            company = entry.get("company", "")
            dates = entry.get("dates", "")
            
            # Create header line with title, company, and dates
            p = doc.add_paragraph()
            _add_section_spacing(p, before=8, after=2)
            
            # Title and company on same line
            header_text = title
            if company:
                header_text += f" | {company}" if header_text else company
            
            run = p.add_run(header_text or "Role")
            _set_font(run, size=11, bold=True)
            
            # Dates on next line (smaller, not bold)
            if dates:
                p = doc.add_paragraph()
                _add_section_spacing(p, before=0, after=4)
                run = p.add_run(dates)
                _set_font(run, size=10)
            
            # Bullet points
            bullets = entry.get("bullets", [])
            logger.debug(f"Adding {len(bullets)} bullets for this experience")
            
            for bullet in bullets:
                bp = doc.add_paragraph(style=None)
                bp.paragraph_format.left_indent = Inches(0.25)
                bp.paragraph_format.first_line_indent = Inches(-0.25)
                _add_section_spacing(bp, before=2, after=2)
                brun = bp.add_run(f"• {bullet}")
                _set_font(brun, size=11)

    # --- Education ---
    education = tailored.get("education") or []
    if education:
        logger.info(f"Adding Education section with {len(education)} entries")
        _add_heading(doc, "Education")
        
        for line in education:
            p = doc.add_paragraph()
            _add_section_spacing(p, before=4, after=2)
            run = p.add_run(line)
            _set_font(run, size=11)

    # --- Certifications ---
    certifications = tailored.get("certifications") or []
    if certifications:
        logger.info(f"Adding Certifications section with {len(certifications)} entries")
        _add_heading(doc, "Certifications")
        
        for line in certifications:
            p = doc.add_paragraph()
            _add_section_spacing(p, before=4, after=2)
            run = p.add_run(line)
            _set_font(run, size=11)

    logger.info(f"Saving document to {output_path}")
    doc.save(output_path)
    logger.info("=== DOCX resume building completed successfully ===")
    return output_path
